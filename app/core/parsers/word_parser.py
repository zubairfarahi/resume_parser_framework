"""Word document parser implementation using python-docx library.

This module provides a concrete implementation of FileParser for Word documents.
"""

import signal
from collections.abc import Generator
from contextlib import contextmanager
from pathlib import Path

from docx import Document

from app.config.logging_config import get_logger, log_performance
from app.core.parsers.base import FileParser
from app.exceptions.exceptions import ParsingError, TimeoutError

logger = get_logger(__name__)


class WordParser(FileParser):
    """Word document (.docx) parser implementation.

    Uses python-docx library to extract text from Word documents with timeout protection.

    SOLID Principles:
    - Single Responsibility: Only handles Word document parsing
    - Liskov Substitution: Can replace FileParser without breaking code
    - Dependency Inversion: Depends on FileParser abstraction

    Example:
        >>> parser = WordParser(timeout=30)
        >>> text = parser.parse(Path("resume.docx"))
    """

    @contextmanager
    def _timeout_handler(self, seconds: int) -> Generator[None, None, None]:
        """Context manager for timeout protection.

        Args:
            seconds: Timeout in seconds

        Raises:
            TimeoutError: If operation exceeds timeout

        Yields:
            None
        """

        def timeout_signal_handler(signum: int, frame: object) -> None:
            raise TimeoutError(
                f"Word parsing exceeded {seconds}s timeout",
                timeout_seconds=seconds,
                operation="word_parsing",
            )

        # Set alarm signal
        old_handler = signal.signal(signal.SIGALRM, timeout_signal_handler)
        signal.alarm(seconds)

        try:
            yield
        finally:
            # Reset alarm
            signal.alarm(0)
            signal.signal(signal.SIGALRM, old_handler)

    @log_performance("word_parsing")
    def parse(self, file_path: Path) -> str:
        """Parse Word document and extract text content.

        Args:
            file_path: Path to Word document

        Returns:
            Extracted text content as string

        Raises:
            FileNotFoundError: If file doesn't exist
            ParsingError: If Word parsing fails
            TimeoutError: If parsing exceeds timeout
        """
        # Validate file first
        validation_error = self.validate_file(file_path)
        if validation_error:
            logger.error("Word validation failed", file_path=str(file_path), error=validation_error)
            raise ParsingError(validation_error, file_path=str(file_path))

        logger.info("Starting Word parsing", file_path=str(file_path))

        try:
            with self._timeout_handler(self.timeout):
                # Read Word document
                document = Document(str(file_path))

                # Extract text from all paragraphs
                text_parts = []

                # Extract from paragraphs
                for paragraph in document.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # Extract from tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)

                # Combine all text
                full_text = "\n".join(text_parts)

                if not full_text.strip():
                    logger.warning("No text extracted from Word document", file_path=str(file_path))
                    raise ParsingError(
                        "No text content found in Word document",
                        file_path=str(file_path),
                    )

                logger.info(
                    "Word parsing successful",
                    file_path=str(file_path),
                    paragraphs=len(document.paragraphs),
                    tables=len(document.tables),
                    text_length=len(full_text),
                )

                return full_text

        except TimeoutError:
            logger.error("Word parsing timeout", file_path=str(file_path), timeout=self.timeout)
            raise

        except Exception as e:
            logger.error(
                "Word parsing failed",
                file_path=str(file_path),
                error=str(e),
                exc_info=True,
            )
            raise ParsingError(
                f"Failed to parse Word document: {str(e)}",
                file_path=str(file_path),
                details={"error_type": type(e).__name__},
            )

    def supports_format(self, file_path: Path) -> bool:
        """Check if this parser supports Word format.

        Args:
            file_path: Path to file

        Returns:
            True if file has .docx extension
        """
        return file_path.suffix.lower() in [".docx", ".doc"]
